########################################
#This program reads CPU information and network configuration details from your current system.
#It stores the results in MS word format
########################################

#import the libraries
import socket
import psutil

#create document object
from docx import Document
document = Document()

#add header
section = document.sections[0]
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = "Laptop Information - LLD"

#add heading for the first page of the document
document.add_heading('System and Network Configuration(Low Level Design Document)', level = 1)
document.add_page_break()

#add heading for the second page
document.add_heading('Introduction')
paragraph = document.add_paragraph('The program reads configuration of the laptop like CPU, network config etc. and display the result')
document.add_page_break()

#Retreive information and update the document

###Basic System Info>>>
document.add_heading('Basic System Info')

hostname = socket.gethostname()   #get host name of the current system 
IPAddr = socket.gethostbyname(hostname) #get IP address of the current system
no_of_phy_cpu = psutil.cpu_count(logical=False) #return no of physical CPUs
no_of_log_cpu = psutil.cpu_count() #return no of Logical CPUs

#create a dict to store these info
dict_system_info = {}
dict_system_info['HostName'] = hostname
dict_system_info['IPAddress'] = IPAddr
dict_system_info['NoOfPhysicalCPU'] = no_of_phy_cpu
dict_system_info['NoOfLogicalCPU'] = no_of_log_cpu

#create a table to store the info
#we have 2 rows and 4 columns
table = document.add_table(rows=2, cols=4)
table.style = 'TableGrid'
hdr_cells = table.rows[0].cells
row_cells = table.rows[1].cells


for i,(key,value) in enumerate(dict_system_info.items()):
    #print(i,key,value)
    hdr_cells[i].text = key
    row_cells[i].text = str(value)

document.add_page_break()
###Basic System Info<<<

###System CPU Frequencies>>>>
document.add_heading('System CPU Frequencies')

#Return CPU frequency as a nameduple including current, min and max frequencies expressed in Mhz
nt_cpu_freq = psutil.cpu_freq(percpu=False)
dict_cpu_freq = nt_cpu_freq._asdict()#convert from named touple to orderdict

#create a table to store the info
#we have 2 rows and 3 columns
table = document.add_table(rows=2, cols=3)
table.style = 'TableGrid'
hdr_cells = table.rows[0].cells
row_cells = table.rows[1].cells


for i,(key,value) in enumerate(dict_cpu_freq.items()):
    #print(i,key,value)
    hdr_cells[i].text = key
    row_cells[i].text = str(value)

document.add_page_break()
###System CPU Frequencies<<<

###System CPU Times>>>>
document.add_heading('System CPU Times')

#add cpu_times as table to the document
#Below function returns system CPU times as a named tuple. Every attribute represents the seconds the CPU has spent in the given mode
nt_cpu_times = psutil.cpu_times()
dict_cpu_times = nt_cpu_times._asdict()#convert from named touple to orderdict

#https://python-docx.readthedocs.io/en/latest/#
#create a table to store the info
#we have 2 rows and 5 columns
table = document.add_table(rows=2, cols=5)
table.style = 'TableGrid'
hdr_cells = table.rows[0].cells
row_cells = table.rows[1].cells


for i,(key,value) in enumerate(dict_cpu_times.items()):
    #print(i,key,value)
    hdr_cells[i].text = key
    row_cells[i].text = str(value)

document.add_page_break()
###System CPU Times<<<

###CPU Memory>>>>
document.add_heading('CPU memory details')

#return total and available memory
nt_cpu_memory = psutil.virtual_memory()
dict_cpu_memory = nt_cpu_memory._asdict()#convert from named touple to orderdict

#create a table to store the info
#we have 2 rows and 5 columns
table = document.add_table(rows=2, cols=5)
table.style = 'TableGrid'
hdr_cells = table.rows[0].cells
row_cells = table.rows[1].cells


for i,(key,value) in enumerate(dict_cpu_memory.items()):
    #print(i,key,value)
    hdr_cells[i].text = key
    row_cells[i].text = str(value)

document.add_page_break()
###CPU Memory<<<

###Network I/O Statistics>>>
document.add_heading('Network I/O Statistics')
#Return system-wide network I/O statistics as a named tuple
nt_net_io_counters = psutil.net_io_counters(pernic=False, nowrap=True)
dict_net_io_counters = nt_net_io_counters._asdict()#convert from named touple to orderdict

#create a table to store the info
#we have 2 rows and 8 columns
table = document.add_table(rows=2, cols=8)
table.style = 'TableGrid'
hdr_cells = table.rows[0].cells
row_cells = table.rows[1].cells


for i,(key,value) in enumerate(dict_net_io_counters.items()):
    #print(i,key,value)
    hdr_cells[i].text = key
    row_cells[i].text = str(value)

document.add_page_break()

###Network I/O Statistics<<<

document.add_heading('References')

#save the document
document.save('SystemInformationLLD.docx')
